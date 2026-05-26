"""把包内手册资源部署到 ``<CHAYUAN_ROOT>/manuals/``,首启时跑一次。

行为
====

* **幂等**:目标文件已存在且大小相同 → 跳过。
* **失败软降级**:python-docx 不可用 → 仅复制 .md;CHAYUAN_ROOT 不可写 →
  记日志但不抛(不阻塞 server 启动)。
* **多手册扩展**:未来增加新手册只需在 :data:`MANUAL_FILES` 注册条目。

部署后路径
==========

::

    <CHAYUAN_ROOT>/manuals/
        察元使用手册.md     ← 来自包内 user_manual.md
        察元使用手册.docx   ← 由 markdown 转换(python-docx)

API
===

* :func:`deploy_user_manuals` —— 入口,首启 hook 调用
* :func:`get_manual_path` —— 按公开名拿磁盘绝对路径
* :func:`list_deployed_manuals` —— 列出已部署条目(给 admin route 用)
"""
from __future__ import annotations

import logging
import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

logger = logging.getLogger("chayuan.manuals.deploy")


# ``<CHAYUAN_ROOT>/<MANUALS_DIRNAME>``
MANUALS_DIRNAME = "manuals"


@dataclass(frozen=True)
class ManualSpec:
    """单本手册的部署描述。"""

    resource: str        # 包内文件名(``user_manual.md``)
    public_name: str     # 部署后中文名(``察元使用手册``)
    description: str = ""

    @property
    def md_filename(self) -> str:
        return f"{self.public_name}.md"

    @property
    def docx_filename(self) -> str:
        return f"{self.public_name}.docx"


#: 当前已注册的手册清单。
MANUAL_FILES: List[ManualSpec] = [
    ManualSpec(
        resource="user_manual.md",
        public_name="察元使用手册",
        description="装机即用功能总览,含快速上手 / 知识库 / 模型管理 / 故障排查",
    ),
    ManualSpec(
        resource="about_chayuan.md",
        public_name="察元产品介绍",
        description="察元 AI 桌面版产品介绍:字词级审核 / 五类知识源 / 六类本地能力 / 两档安装包",
    ),
]


# ───────────────────────── 资源访问 ─────────────────────────


def _resources_dir() -> Path:
    """包内 resources 目录绝对路径。"""
    return Path(__file__).resolve().parent / "resources"


def _read_resource(name: str) -> str:
    path = _resources_dir() / name
    return path.read_text(encoding="utf-8")


def _manuals_dir() -> Optional[Path]:
    """``<CHAYUAN_ROOT>/manuals``。CHAYUAN_ROOT 不可用时返回 None。"""
    try:
        from chayuan.settings import CHAYUAN_ROOT
    except Exception as e:  # noqa: BLE001
        logger.debug("[manuals] CHAYUAN_ROOT 不可用: %r", e)
        return None
    return Path(CHAYUAN_ROOT) / MANUALS_DIRNAME


# ───────────────────────── Markdown → docx ─────────────────────────


def _markdown_to_docx(md_text: str, out_path: Path) -> bool:
    """用 python-docx 把 markdown 转成 .docx。成功 True,失败 False(不抛)。

    转换规则(简化版,够手册看):
    * ``# / ## / ### / ####`` → 对应 Heading level
    * ```` ``` `` 围栏代码块 → Monospace 段落
    * ``| col | col |`` 表格行(连续) → docx 表格
    * ``- item`` / ``* item`` / ``1. item`` → ListBullet / ListNumber
    * ``> quote`` → 引用样式
    * 段落 → 普通段落

    不做行内格式(粗体 / 链接 / 行内代码),手册阅读已经够用;追求富文本
    可以打开 .md 原文。
    """
    try:
        from docx import Document  # type: ignore[import-not-found]
    except Exception as e:  # noqa: BLE001
        logger.debug("[manuals] python-docx 不可用,跳过 docx 生成: %r", e)
        return False

    try:
        doc = Document()
        lines = md_text.splitlines()
        i = 0
        in_code = False
        code_buf: list[str] = []
        table_buf: list[list[str]] = []

        def _flush_table() -> None:
            if not table_buf:
                return
            cols = max(len(row) for row in table_buf)
            t = doc.add_table(rows=len(table_buf), cols=cols)
            t.style = "Light Grid Accent 1"
            for ri, row in enumerate(table_buf):
                for ci in range(cols):
                    cell = t.rows[ri].cells[ci]
                    cell.text = row[ci] if ci < len(row) else ""
            table_buf.clear()

        while i < len(lines):
            line = lines[i]

            # 围栏代码块
            if line.lstrip().startswith("```"):
                if in_code:
                    p = doc.add_paragraph("\n".join(code_buf))
                    p.style = "Intense Quote"
                    code_buf.clear()
                    in_code = False
                else:
                    _flush_table()
                    in_code = True
                i += 1
                continue
            if in_code:
                code_buf.append(line)
                i += 1
                continue

            # 表格
            if "|" in line and re.match(r"^\s*\|.*\|\s*$", line):
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                # 分隔行 `|---|---|` 跳过
                if all(re.match(r"^:?-+:?$", c) for c in cells if c):
                    i += 1
                    continue
                table_buf.append(cells)
                i += 1
                continue
            else:
                _flush_table()

            # 标题
            m = re.match(r"^(#{1,6})\s+(.*)$", line)
            if m:
                level = len(m.group(1))
                doc.add_heading(m.group(2), level=min(level, 4))
                i += 1
                continue

            # 引用
            if line.startswith(">"):
                p = doc.add_paragraph(line.lstrip("> ").rstrip())
                p.style = "Intense Quote"
                i += 1
                continue

            # 列表
            if re.match(r"^\s*[-*]\s+", line):
                content = re.sub(r"^\s*[-*]\s+", "", line)
                doc.add_paragraph(content, style="List Bullet")
                i += 1
                continue
            if re.match(r"^\s*\d+\.\s+", line):
                content = re.sub(r"^\s*\d+\.\s+", "", line)
                doc.add_paragraph(content, style="List Number")
                i += 1
                continue

            # 普通段落
            if line.strip():
                doc.add_paragraph(line)
            i += 1

        _flush_table()
        if in_code and code_buf:
            p = doc.add_paragraph("\n".join(code_buf))
            p.style = "Intense Quote"

        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return True
    except Exception as e:  # noqa: BLE001
        logger.warning("[manuals] markdown_to_docx 异常: %r", e)
        return False


# ───────────────────────── 部署主入口 ─────────────────────────


@dataclass
class DeployReport:
    """:func:`deploy_user_manuals` 返回的报告。"""

    target_dir: Optional[str]
    md_written: List[str] = field(default_factory=list)   # public_name 列表
    docx_written: List[str] = field(default_factory=list)
    skipped: List[str] = field(default_factory=list)      # 已是最新,跳过
    errors: List[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "target_dir": self.target_dir,
            "md_written": list(self.md_written),
            "docx_written": list(self.docx_written),
            "skipped": list(self.skipped),
            "errors": list(self.errors),
        }


def deploy_user_manuals(*, force: bool = False) -> DeployReport:
    """把包内手册资源复制到 ``<CHAYUAN_ROOT>/manuals/`` 并生成 docx。

    Args:
        force: ``True`` 时覆盖已存在文件(忽略内容比对)。默认幂等。

    Returns:
        :class:`DeployReport`。
    """
    target = _manuals_dir()
    report = DeployReport(target_dir=str(target) if target else None)
    if target is None:
        report.errors.append("CHAYUAN_ROOT 不可用,跳过部署")
        return report

    try:
        target.mkdir(parents=True, exist_ok=True)
    except OSError as e:
        report.errors.append(f"创建 {target} 失败: {e}")
        return report

    for spec in MANUAL_FILES:
        try:
            md_src = _read_resource(spec.resource)
        except FileNotFoundError as e:
            report.errors.append(f"包内资源缺失: {spec.resource} ({e})")
            continue

        md_path = target / spec.md_filename
        # 内容对比:已是最新则跳过(支持手册版本升级覆盖)
        if md_path.is_file() and not force:
            try:
                if md_path.read_text(encoding="utf-8") == md_src:
                    report.skipped.append(spec.public_name)
                    # docx 也可能已经存在;若不存在仍尝试生成
                    docx_path = target / spec.docx_filename
                    if not docx_path.is_file():
                        if _markdown_to_docx(md_src, docx_path):
                            report.docx_written.append(spec.public_name)
                    continue
            except OSError as e:
                report.errors.append(f"读取 {md_path} 失败: {e}")
                continue

        try:
            md_path.write_text(md_src, encoding="utf-8")
            report.md_written.append(spec.public_name)
        except OSError as e:
            report.errors.append(f"写入 {md_path} 失败: {e}")
            continue

        docx_path = target / spec.docx_filename
        if _markdown_to_docx(md_src, docx_path):
            report.docx_written.append(spec.public_name)

    logger.info(
        "[manuals] deploy_user_manuals: md=%d docx=%d skipped=%d errors=%d → %s",
        len(report.md_written), len(report.docx_written),
        len(report.skipped), len(report.errors), report.target_dir,
    )
    return report


# ───────────────────────── 查询入口 ─────────────────────────


def list_deployed_manuals() -> List[Dict[str, object]]:
    """列出已部署的手册(给 admin route 用)。

    每条 dict 含 ``public_name`` / ``description`` / ``md_path`` /
    ``docx_path`` / ``md_exists`` / ``docx_exists``。
    """
    target = _manuals_dir()
    out: List[Dict[str, object]] = []
    for spec in MANUAL_FILES:
        md_path = (target / spec.md_filename) if target else None
        docx_path = (target / spec.docx_filename) if target else None
        out.append({
            "public_name": spec.public_name,
            "description": spec.description,
            "md_path":   str(md_path) if md_path else None,
            "docx_path": str(docx_path) if docx_path else None,
            "md_exists":   bool(md_path and md_path.is_file()),
            "docx_exists": bool(docx_path and docx_path.is_file()),
        })
    return out


def get_manual_path(public_name: str, *, fmt: str = "docx") -> Optional[Path]:
    """按公开名拿磁盘文件路径;不存在返回 None。

    Args:
        public_name: 例如 "察元使用手册"
        fmt: ``docx`` 或 ``md``
    """
    target = _manuals_dir()
    if target is None:
        return None
    for spec in MANUAL_FILES:
        if spec.public_name != public_name:
            continue
        path = target / (spec.docx_filename if fmt == "docx" else spec.md_filename)
        return path if path.is_file() else None
    return None


__all__ = [
    "DeployReport",
    "ManualSpec",
    "MANUALS_DIRNAME",
    "MANUAL_FILES",
    "deploy_user_manuals",
    "get_manual_path",
    "list_deployed_manuals",
]
