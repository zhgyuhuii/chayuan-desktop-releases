from typing import List, Optional

import tqdm
from langchain_community.document_loaders.unstructured import UnstructuredFileLoader

from chayuan.server.file_rag.document_loaders.partition import (
    IMAGE_OCR_PREFIX,
    partition_ocr_text,
)


def _docx_xml_text(filepath: str) -> str:
    """Best-effort text extraction for docx when python-docx cannot parse it."""
    import zipfile
    import xml.etree.ElementTree as ET

    parts: List[str] = []
    with zipfile.ZipFile(filepath) as zf:
        names = [
            name for name in zf.namelist()
            if name == "word/document.xml"
            or name.startswith("word/header")
            or name.startswith("word/footer")
        ]
        for name in names:
            try:
                root = ET.fromstring(zf.read(name))
            except Exception:
                continue
            texts = [
                node.text or ""
                for node in root.iter()
                if node.tag.endswith("}t") and (node.text or "").strip()
            ]
            if texts:
                parts.append("\n".join(texts))
    return "\n\n".join(parts)


def _image_part_class():
    try:
        from docx.parts.image import ImagePart
        return ImagePart
    except Exception:
        try:
            from docx import ImagePart  # type: ignore
            return ImagePart
        except Exception:
            return None


def _build_ocr():
    try:
        from rapidocr_onnxruntime import RapidOCR
        return RapidOCR()
    except Exception:
        return None


def _ocr_image_blob(ocr, blob: Optional[bytes]) -> str:
    if ocr is None or not blob:
        return ""
    try:
        from io import BytesIO

        import numpy as np
        from PIL import Image

        image = Image.open(BytesIO(blob))
        result, _ = ocr(np.array(image))
        if not result:
            return ""
        return "\n".join(line[1] for line in result if len(line) > 1 and line[1])
    except Exception:
        return ""


class RapidOCRDocLoader(UnstructuredFileLoader):
    def _get_elements(self) -> List:
        def doc2text(filepath):
            from docx import Document
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import Table, _Cell
            from docx.text.paragraph import Paragraph

            ocr = _build_ocr()
            image_part_cls = _image_part_class()
            doc = Document(filepath)
            resp = ""

            def iter_block_items(parent):
                from docx.document import Document

                if isinstance(parent, Document):
                    parent_elm = parent.element.body
                elif isinstance(parent, _Cell):
                    parent_elm = parent._tc
                else:
                    raise ValueError("RapidOCRDocLoader parse fail")

                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            b_unit = tqdm.tqdm(
                total=len(doc.paragraphs) + len(doc.tables),
                desc="RapidOCRDocLoader block index: 0",
            )
            for i, block in enumerate(iter_block_items(doc)):
                b_unit.set_description("RapidOCRDocLoader  block index: {}".format(i))
                b_unit.refresh()
                if isinstance(block, Paragraph):
                    resp += block.text.strip() + "\n"
                    images = block._element.xpath(".//pic:pic")  # 获取所有图片
                    for image in images:
                        for img_id in image.xpath(".//a:blip/@r:embed"):  # 获取图片id
                            part = doc.part.related_parts[
                                img_id
                            ]  # 根据图片id获取对应的图片
                            if image_part_cls is None or isinstance(part, image_part_cls):
                                blob = getattr(part, "blob", None) or getattr(part, "_blob", None)
                                ocr_text = _ocr_image_blob(ocr, blob)
                                if ocr_text:
                                    # 来源标注:文档内嵌图片 OCR 出的文字,加前缀
                                    # 并用空行单独成段(不混进正文段落)
                                    resp += f"\n\n{IMAGE_OCR_PREFIX}{ocr_text}\n\n"
                elif isinstance(block, Table):
                    for row in block.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                resp += paragraph.text.strip() + "\n"
                b_unit.update(1)
            b_unit.close()
            return resp

        try:
            text = doc2text(self.file_path)
        except Exception:
            text = _docx_xml_text(self.file_path)
        if not (text or "").strip():
            text = _docx_xml_text(self.file_path)
        return partition_ocr_text(text=text, unstructured_kwargs=self.unstructured_kwargs)


if __name__ == "__main__":
    loader = RapidOCRDocLoader(file_path="../tests/samples/ocr_test.docx")
    docs = loader.load()
    print(docs)
